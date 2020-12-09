import os
import logging
from io import BytesIO
from urllib.request import urlopen
import requests
import datetime

from django.http import HttpResponse

from api.models import Image
from api.utils import set_function_attributes

from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from docxcompose.composer import Composer

import PIL

LOGGER = logging.getLogger('django')

CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
DOC_RESOURCES_PATH = os.path.join(CURRENT_DIR, "..", "..", "..", "doc_resources")
SEAL_IMAGE_PATH = os.path.join(DOC_RESOURCES_PATH, "seal.png")

FIND_TAIWAN_LEGISLATOR_API = "https://ftl.disfactory.tw"
DEFAULT_FONT = "標楷體"

UPPER_CASE_NUMBERS = {
    str(i): chinese_char
    for i, chinese_char in enumerate('零壹貳參肆伍陸柒捌玖')
}
LOWER_CASE_NUMBERS = {
    str(i): chinese_char
    for i, chinese_char in enumerate('〇一二三四五六七八九')
}


def to_lower_chinese_numbers(number):
    return ''.join(LOWER_CASE_NUMBERS[ch] for ch in str(number))


def find_taiwan_legislator_name_by_location(lat, lng):
    try:
        resp = requests.get(
            FIND_TAIWAN_LEGISLATOR_API,
            params={"lat": lat, "lng": lng},
        )
        data = resp.json()
        return data[0]['name'] if (data and 'name' in data[0]) else "UNKNOWN"

    except Exception as e:
        LOGGER.error(
            f"Can't get the legislator information from {FIND_TAIWAN_LEGISLATOR_API} - {e}",
        )
        return "UNKNOWN"


class Run:
    def __init__(self, context, size):
        self._context = context
        self._size = size

    def new(self, paragraph):
        run = paragraph.add_run(self._context)
        run.font.size = Pt(self._size)
        return run


class ParagraphGenerator:
    ALIGN_LEFT = WD_ALIGN_PARAGRAPH.LEFT
    ALIGN_CENTER = WD_ALIGN_PARAGRAPH.CENTER
    ALIGN_RIGHT = WD_ALIGN_PARAGRAPH.RIGHT

    def __init__(self):
        self._alignment = None
        self._line_spacing = None
        self._space_before = None
        self._space_after = None
        self._first_line_indent = None
        self._left_indent = None

    def new(self, document, context, font_size):
        paragraph = document.add_paragraph()
        if self._alignment is not None:
            paragraph.paragraph_format.alignment = self._alignment

        if self._line_spacing is not None:
            paragraph.paragraph_format.line_spacing = self._line_spacing

        if self._space_before is not None:
            paragraph.paragraph_format.space_before = self._space_before

        if self._space_after is not None:
            paragraph.paragraph_format.space_after = self._space_after

        if self._left_indent is not None:
            paragraph.paragraph_format.left_indent = self._left_indent

        if self._first_line_indent is not None:
            paragraph.paragraph_format.first_line_indent = self._first_line_indent

        run = Run(context, font_size)
        run.new(paragraph)

        return paragraph

    def alignment(self, alignment):
        self._alignment = alignment
        return self

    def line_spacing(self, size):
        self._line_spacing = Pt(size)
        return self

    def space_after(self, size):
        self._space_after = Pt(size)
        return self

    def space_before(self, size):
        self._space_before = Pt(size)
        return self

    def left_indent(self, inches_size):
        self._left_indent = Inches(inches_size)
        return self

    def first_line_indent(self, inches_size):
        self._first_line_indent = Inches(inches_size)
        return self


class FactoryReportDocumentWriter:
    def __init__(self, model, document):
        self.document_model = model
        self.factory = model.factory
        self.document = document
        self.factory_location = f"{self.factory.townname}{self.factory.sectname} ({self.factory.sectcode}) {self.factory.landcode}"

        self._generate_docx()

    def _init_document(self):
        # Change layout
        section = self.document.sections[0]

        # Margins
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # A4 paper
        section.page_height = Mm(297)
        section.page_width = Mm(210)

    def _generate_docx(self):
        self._init_document()
        legislator_name = find_taiwan_legislator_name_by_location(
            lat=self.factory.lat,
            lng=self.factory.lng,
        )

        # Cover
        self._original()
        self._title()
        self._sender(self.document_model.cet_staff)
        self._receiver(self.document_model.code)
        self._subject()
        self._context()
        self._cc(legislator_name)
        self._seal()

        self._page_break()

        # Attachments
        self._attchments()

    def _title(self):
        generator = ParagraphGenerator() \
            .alignment(ParagraphGenerator.ALIGN_CENTER) \
            .line_spacing(20) \
            .space_before(6) \
            .space_after(0)

        generator.new(self.document, "地球公民基金會 函", 20)
        generator.new(self.document, "", 20)

    def _original(self):
        generator = ParagraphGenerator() \
            .alignment(ParagraphGenerator.ALIGN_LEFT) \
            .line_spacing(20) \
            .space_before(0) \
            .space_after(0)

        generator.new(self.document, "正本", 12)
        generator.new(self.document, "", 20)

    def _sender(self, sender_name):
        # yapf: disable
        context = [
            '地址：10049台北市北平東路28號9樓之2',
            '電話：02-23920371',
            '傳真：02-23920381',
            '連絡人：{}'.format(sender_name),
            '電子信箱：eva@cet-taiwan.org'
        ]
        # yapf: enable

        generator = ParagraphGenerator() \
            .alignment(ParagraphGenerator.ALIGN_RIGHT) \
            .line_spacing(10) \
            .space_after(0)

        for line in context:
            generator.new(self.document, line, 10)

    def _receiver(self, serial):
        utcnow = datetime.datetime.utcnow().replace(tzinfo=datetime.timezone.utc)
        twnow = utcnow.astimezone(datetime.timezone(datetime.timedelta(hours=8)))

        # yapf: disable
        context = [
            '',
            '受文者：如正、副本行文單位',
            f'發文日期：中華民國{twnow.year}年{twnow.month}月{twnow.day}日',
            f'發文字號：地球公民違字第 {serial} 號',
            '速別：普通件',
            '附件：舉證照片',
            '',
        ]
        # yapf: enable

        generator = ParagraphGenerator() \
            .alignment(ParagraphGenerator.ALIGN_LEFT) \
            .line_spacing(21) \
            .space_after(0)

        for line in context:
            generator.new(self.document, line, 14)

    def _subject(self):
        # yapf: disable
        context = [
            f'主旨：舉報 {self.factory_location} 地號土地疑有違法新增建築情事。',
            '',
            '說明：'
        ]
        # yapf: enable

        generator = ParagraphGenerator() \
            .alignment(ParagraphGenerator.ALIGN_LEFT) \
            .line_spacing(21) \
            .space_after(0)

        for line in context:
            generator.new(self.document, line, 14)

    def _context(self):
        # yapf: disable
        context = [
            '一、　依工廠管理輔導法第28-1、28-12條辦理。',
            f"二、　{self.factory_location} 地號土地新發現新增建情形，經地球公民基金會志工拍攝存證，如附件一。因懷疑係屬非法建築行為，函請貴府調查處理。若有不法情事，並應依法裁處，請貴府將查處情形，惠知本會。"
        ]
        # yapf: enable

        generator = ParagraphGenerator() \
            .alignment(ParagraphGenerator.ALIGN_LEFT) \
            .line_spacing(21) \
            .space_after(0)

        for line in context:
            generator.new(self.document, line, 14)

    def _cc(self, legislator):
        if self.factory.townname:
            townname = self.factory.townname[:3]
        else:
            townname = "UNKNOWN"

        # yapf: enable
        context = [
            '',
            f"正本：{townname}政府",
            f"副本：內政部、行政院農委會、經濟部工業局、經濟部中部辦公室、立法委員{legislator}國會辦公室",
        ]
        # yapf: disable

        generator = ParagraphGenerator() \
            .alignment(ParagraphGenerator.ALIGN_LEFT) \
            .line_spacing(12) \
            .space_after(0)

        for line in context:
            generator.new(self.document, line, 12)

    def _seal(self):
        self.document.add_picture(SEAL_IMAGE_PATH, width=Inches(4.5))
        paragraph = self.document.paragraphs[-1]
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_before = Inches(0.7)

    def _attchments(self):
        generator = ParagraphGenerator() \
            .alignment(ParagraphGenerator.ALIGN_LEFT) \
            .line_spacing(18) \
            .space_after(6)

        images = Image.objects.only("id").filter(factory=self.factory)
        for index, image in enumerate(images, start=1):
            generator.new(self.document, f"附件 {to_lower_chinese_numbers(index)}", 12)
            data = urlopen(image.image_path).read()

            image_data = PIL.Image.open(BytesIO(data))
            if image_data.format == "JPEG":
                # Use PIL to save all jpeg files again to workaround python-docx bug
                # https://github.com/python-openxml/python-docx/issues/187
                tmp_image_data = BytesIO()
                image_data.save(tmp_image_data, format="jpeg")
                self.document.add_picture(tmp_image_data, width=Mm(150))
            else:
                self.document.add_picture(BytesIO(data), width=Mm(150))

    def _page_break(self):
        self.document.add_page_break()


def new_document():
    document = Document()
    document.styles['Normal'].font.name = DEFAULT_FONT
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)
    return document


def generate_factories_document(model_list):
    documents = []
    for model in model_list:
        document = new_document()
        FactoryReportDocumentWriter(model, document)
        documents.append(document)
    return documents


def merge_documents(documents):
    for document in documents[:-1]:
        document.add_page_break()

    composer = Composer(documents[0])
    for document in documents[1:]:
        composer.append(document)

    return composer.doc


def export_document(document):
    file = BytesIO()
    document.save(file)

    return file


class ExportDocMixin:

    @set_function_attributes(short_description="輸出成 docx 檔")
    def export_as_docx(self, request, queryset):
        documents = generate_factories_document(list(queryset))
        merged_docment = merge_documents(documents)
        docx_file = export_document(merged_docment)

        length = docx_file.tell()
        docx_file.seek(0)

        response = HttpResponse(
            docx_file.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = 'attachment; filename=api.factory.docx'
        response['Content-Length'] = length
        return response
