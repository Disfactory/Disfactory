import csv
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING 
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import copy
from io import BytesIO

def _modify(field_name):
    if field_name == 'get_name':
        return 'name'
    return field_name


class ExportCsvMixin:

    def export_as_csv(self, request, queryset):
        meta = self.model._meta
        field_names = [_modify(field) for field in self.list_display]

        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename={}.csv'.format(meta)
        response.write(u'\ufeff'.encode('utf8'))
        writer = csv.writer(response)

        writer.writerow(field_names)
        for obj in queryset:
            writer.writerow([getattr(obj, field) for field in field_names])

        return response

    export_as_csv.short_description = '輸出成 csv 檔'

class ExportDocMixin:

    def export_as_docx(self, request, queryset):
        meta = self.model._meta
        field_names = [_modify(field) for field in self.list_display]

        doclist = []
        for idx, obj in enumerate(queryset):
            if idx < len(queryset)-1:
                doclist.append(self.CreateDoc(obj, 20200701,page_break=True))
            else:
                doclist.append(self.CreateDoc(obj, 20200701,page_break=False))
        #doclist.append(self.CreateDoc(456))
        #doclist.append(self.CreateDoc(948794))
        
        docs = []
        for idx, item in enumerate(doclist):
            if idx > 0:
                #docs[idx].add_page_break()
                for element in item.element.body:
                    docs[0].element.body.append(element)
            else:
                docs.append(item)        
    
        f = BytesIO()
        docs[0].save(f)
        length = f.tell()
        f.seek(0)
        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=' + 'api.factory.docx'
        response['Content-Length'] = length
        return response
    
    export_as_docx.short_description = '輸出成 docx 檔'

    def ChineseNumber(number, flag='Minor'):
        Major = {0: '零', 1: '壹', 2: '貳', 3: '參', 4: '肆',
                   5: '伍', 6: '陸', 7: '柒', 8: '捌', 9: '玖'}
        Minor = {0: '〇', 1: '一', 2: '二', 3: '三', 4: '四',
                   5: '五', 6: '六', 7: '七', 8: '八', 9: '九'}
        if flag == 'Minor':
            return Minor[number]
        else:
            return Major[number]
        
    def ChineseNumberList(number, flag='Minor'):
        # Returns a string of chinese formal number from number
        numberchar = []
        numberbase = 1
        while number >= numberbase:
            numberchar += ChineseNumber(number % (numberbase * 10) // numberbase, flag)
            numberbase *= 10
        return ''.join(numberchar[::-1])
    
    def CreateDoc(self, obj, serial, page_break=False):
        document = Document('./doc_templates/doc_template_element_only.docx')
        #document = Document()
        # Change layout
        section = document.sections[0]
        # Margins
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        # A4 paper
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        
        # Title
        

        p0 = ['地球公民基金會 函',
              '']
        
        for p in p0:
            par = document.add_paragraph()
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.line_spacing = Pt(20)
            par.paragraph_format.space_before = Pt(6)
            par.paragraph_format.space_after = 0
            run = par.add_run(p)
            run.font.name = 'BiauKai'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'BiauKai')
            run.font.size = Pt(20)
        
        
        
        # Contact - chagne based on who's sending it
        
        p1 = ['地址：10049台北市北平東路28號9樓之2',
              '電話：02-23920371 ',
              '傳真：02-23920381',
              '連絡人：吳沅諭',
              '電子信箱：eva@cet-taiwan.org'
              ]
        
        for p in p1:
            par = document.add_paragraph()
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.line_spacing = Pt(12)
            par.paragraph_format.space_after = 0
            par.paragraph_format.left_indent = Inches(3.6)
            run = par.add_run(p)
            run.font.name = 'BiauKai'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'BiauKai')
            run.font.size = Pt(10)
        
        
        p2 = ['',
              '受文者：如正、副本行文單位']
        for p in p2:
            par = document.add_paragraph()
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.line_spacing = 1
            par.paragraph_format.space_after = 0
            run = par.add_run(p)
            run.font.name = 'BiauKai'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'BiauKai')
            run.font.size = Pt(14)
            
            
        # Change time and serial number as needed
        
        p3 = ['發文日期：中華民國109年6月31日',
              f'發文字號：地球公民違字第 {serial} 號',
              '速別：普通件',
              '附件：舉證照片',
              '']
        for p in p3:
            par = document.add_paragraph()
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.line_spacing = 1
            par.paragraph_format.space_after = 0
            run = par.add_run(p)
            run.font.name = 'BiauKai'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'BiauKai')
            run.font.size = Pt(12)
        
        
        # Change location as needed
        
        p4 = [f'主旨：舉報{getattr(obj, "townname")}({getattr(obj, "sectcode")}) {getattr(obj, "landcode")} 地號土地疑有違法新增建築情事。',
            '說明：']
        
        for p in p4:
            par = document.add_paragraph()
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.line_spacing = Pt(25)
            par.paragraph_format.space_after = 0
            par.paragraph_format.first_line_indent = Inches(-0.56)
            par.paragraph_format.left_indent = Inches(0.56)
            run = par.add_run(p)
            run.font.name = 'BiauKai'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'BiauKai')
            run.font.size = Pt(14)
        
        
        # Change location again as needed
        
        p5 = ['一、　依工廠管理輔導法第28-1、28-12條辦理。',
              ''
              f'二、　{getattr(obj, "townname")}({getattr(obj, "sectcode")}) {getattr(obj, "landcode")} 地號土地新發現新增建情形，經地球公民基金會志工拍攝存證，如附件一。因懷疑係屬非法建築行為，函請貴府調查處理。若有不法情事，並應依法裁處，請貴府將查處情形，惠知本會。'
              ]
        
        for p in p5:
            par = document.add_paragraph()
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.line_spacing = 1
            par.paragraph_format.space_after = 0
            run = par.add_run(p)
            run.font.name = 'BiauKai'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'BiauKai')
            run.font.size = Pt(14)
        
        
        # Change addressee and legislator
        
        p6 =['',
             f'正本：{getattr(obj,"townname")[:3]}政府',
             '副本：內政部、行政院農委會、經濟部工業局、經濟部中部辦公室、立法委員沒有人國會辦公室'
             ]
        
        for p in p6:
            par = document.add_paragraph()
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.line_spacing = 1
            par.paragraph_format.space_after = 0
            run = par.add_run(p)
            run.font.name = 'BiauKai'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'BiauKai')
            run.font.size = Pt(12)
        
        document.add_picture('./doc_templates/seal.png', width=Inches(4.5))
        par = document.paragraphs[-1]
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        par.paragraph_format.space_before = Inches(0.7)
        
        # List attachments in the next pages
        if False:
            document.add_page_break()
            
            # Given a list of images, list them one by one
             
            images = ['Picture1.png', 'Picture1.png']
            
            for idx, img in enumerate(images):
                par = document.add_paragraph()
                par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                par.paragraph_format.line_spacing = 1
                par.paragraph_format.space_after = Pt(6)
                run = par.add_run('附件' + ChineseNumberList(idx+1))
                run.font.name = 'BiauKai'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'BiauKai')
                run.font.size = Pt(12)
                document.add_picture(img, width=Inches(4))
                par = document.add_paragraph() # For a new line
        if True: 
            if page_break:
                document.add_page_break()
    
        return document

class RestoreMixin:

    def restore(self, request, queryset):
        queryset.undelete()

    restore.short_description = '復原'


class ExportLabelMixin:
    def export_labels_as_docx(self, request, queryset):
        return

    export_labels_as_docx.short_description = '下載標籤及交寄執據'

